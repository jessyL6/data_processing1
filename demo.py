# -*-coding:utf-8 -*-
import jieba
import docx  # 处理docx文档
from collections import Counter
from pdfminer.converter import PDFPageAggregator  # 处理pdf
from pdfminer.layout import LAParams
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice


def get_txt():
    f1 = open("莲.txt", "r", encoding='cp936')
    f2 = open("莲-词语统计.txt", "w", encoding='cp936')
    txt = f1.read()
    words = jieba.cut(txt)

    #处理停用词
    words = handle(words)

    # 方法一：使用计数器
    c = Counter()
    for x in words:
        c[x] += 1
    print('常用词频度统计结果')
    ls = []
    # most_commen是提取前n个最常用的
    for (k, v) in c.most_common(100):
        print('%s %d' % (k, v))
        #print(type((k, v)))
        # 类型是元组
        ls.append(v)
        d = k + ':' + str(v) + '\n'
        f2.write(d)
    print(ls)
    f2.write(str(ls))
    f1.close()
    f2.close()


def get_doc():
    # python读取docx文档需要使用docx库
    f = docx.Document("八百标兵.docx")
    f2 = docx.Document()

    # f.paragraphs是段落
    # print("段落数:" + str(len(f.paragraphs)))
    text = ""  # 接收word文本内容
    for para in f.paragraphs:
        text += para.text
    # print(text)
    # 对过滤后的文本进行分词
    words = jieba.cut(text, cut_all=False)

    words = handle(words)

    # 统计词频
    word_freq = {}  # 词频序列
    for word in words:
        if word in word_freq:  # 统计字/词出现的次数，即若是该词已存在于序列中，则次数+1
            word_freq[word] += 1
        else:  # 若是不存在，那么把该词加入序列中
            word_freq[word] = 1
    freq_word = []  # 方法二： 排序字/词列表,lambla排序
    # 向量矩阵
    ls = []
    for word, freq in word_freq.items():
        freq_word.append((word, freq))
    freq_word.sort(key=lambda x: x[1], reverse=True)
    for word2, freq2 in freq_word:
        ls.append(freq2)
        f2.add_paragraph(str((word2, freq2)))
    print(freq_word)
    print(ls)
    f2.add_paragraph(str(ls))
    f2.save("八百标兵_字词统计.docx")


def get_pdf():
    f = open("pdf结果.txt", "w", encoding='cp936')
    # 基本思路：将pdf转换为成txt
    text = transport()
    # 除掉多余的空格和换行符
    text = text.replace("\n", "")
    text = text.replace(" ", "")
    # print(text)
    print("yes")
    words = jieba.cut(text, cut_all=False)

    #去除停用词
    words = handle(words)

    # 统计词频
    word_freq = {}  # 词频序列
    for word in words:
        if word in word_freq:  # 统计字/词出现的次数，即若是该词已存在于序列中，则次数+1
            word_freq[word] += 1
        else:  # 若是不存在，那么把该词加入序列中
            word_freq[word] = 1
    freq_word = []  # 方法二： 排序字/词列表,lambla排序
    # 向量矩阵
    ls = []
    for word, freq in word_freq.items():
        freq_word.append((word, freq))
    freq_word.sort(key=lambda x: x[1], reverse=True)

    #print(type(freq_word))
    print(freq_word)
    for word2, freq2 in freq_word:
        ls.append(freq2)
        f.write(str((word2, freq2)) + '\n')
    # print(freq_word)
    # print(ls)
    f.write(str(ls))

def handle(words):
    # 去除自定义停用词，使用自己定义的函数,但是在这里进行去除并且不再次jieba.cut的话，会产生分出的词为单独的一个个的字
    words = movestopwords(words)

    # 在这里需要再次jieba.cut分词，不然会出现上面的情况
    words = jieba.cut(words, cut_all=False)
    return words

def transport():
    # 获取文档对象
    fp = open("新增80道二级公共基础选择题.pdf", "rb")

    # 创建一个一个与文档关联的解释器
    parser = PDFParser(fp)

    # PDF文档的对象
    doc = PDFDocument()

    # 连接解释器和文档对象
    parser.set_document(doc)
    doc.set_parser(parser)

    # 初始化文档,当前文档没有密码，设为空字符串
    doc.initialize("")

    # 创建PDF资源管理器
    resource = PDFResourceManager()

    # 参数分析器
    laparam = LAParams()

    # 创建一个聚合器
    device = PDFPageAggregator(resource, laparams=laparam)

    # 创建PDF页面解释器
    interpreter = PDFPageInterpreter(resource, device)

    ls2 = []
    # 使用文档对象得到页面的集合
    for page in doc.get_pages():
        # 使用页面解释器读取
        interpreter.process_page(page)

        # 使用聚合器来获得内容
        layout = device.get_result()

        for out in layout:
            if hasattr(out, "get_text"):
                # print(out.get_text())
                ls2.append(out.get_text())
    ls2 = ''.join(ls2)
    # print(ls2)
    return ls2


# 创建停用词列表，从自己写的停用词txt文件中读取出来
def stopwordslist(fpath):
    stopwords = [line.strip() for line in open(fpath, 'r', encoding='cp936')]
    return stopwords


# 除去停用词
# 这里需要修改，因为传参过来是元组，字典？？？？
def movestopwords(sentence):
    stopwords = stopwordslist('stopwords.txt')  # 这里加载停用词的路径
    #print(stopwords)
    outstr = ''
    for word in sentence:
        if word not in stopwords:
            outstr += word
            # outstr += " "
    #print(outstr)
    #注意返回的是字符串
    return outstr


def main():
    get_txt()
    get_doc()
    get_pdf()
    print("完成")


if __name__ == '__main__':
    main()
